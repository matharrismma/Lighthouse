"""Extract devotional reflections + sermons from Matt's iCloud .docx files
into structured JSONL substrate for the Concordance engine.

One-shot script — run once. Output:
  data/devotionals/reflections.jsonl
  data/sermons/sermons.jsonl
"""
from __future__ import annotations
import json
import re
import sys
import hashlib
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx not installed. pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

REPO = Path(__file__).parent.parent
SOURCE_DIR = Path.home() / "iCloudDrive" / "Christian writing"
DEVOTIONALS_OUT = REPO / "data" / "devotionals" / "reflections.jsonl"
SERMONS_OUT = REPO / "data" / "sermons" / "sermons.jsonl"

DEVOTIONAL_BATCHES = [
    SOURCE_DIR / f"Devotional_Reflections_Chronological_Batch{i}.docx" for i in range(1, 7)
]

# US date format mm/dd/yyyy from the source files
DATE_RE = re.compile(r"^Date:\s*(\d{1,2})/(\d{1,2})/(\d{4})\s*$")
SCRIPTURE_HEAD_RE = re.compile(r"^Scripture\s*(?:\(([^)]+)\))?:\s*$")


def _slugify(text: str, maxlen: int = 40) -> str:
    s = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return s[:maxlen] or "untitled"


def _short_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:6]


def extract_devotional_entries(path: Path) -> list[dict]:
    """Parse one devotional batch .docx into a list of structured entries.

    Anchor on Date: lines. Each Date: line marks one entry. The title is
    the nearest preceding non-skip, non-Scripture paragraph that is not
    itself a Date: line. Body + Scripture extend forward until the next
    Date: line.
    """
    d = Document(path)
    paragraphs = [p.text.strip() for p in d.paragraphs]
    entries: list[dict] = []

    skip_titles = {"Devotional Reflections Collection", "Table of Contents", ""}

    # Find every Date: line and its position
    date_positions = [i for i, p in enumerate(paragraphs) if DATE_RE.match(p)]

    for idx, date_pos in enumerate(date_positions):
        m = DATE_RE.match(paragraphs[date_pos])
        month, day, year = map(int, m.groups())
        iso_date = f"{year:04d}-{month:02d}-{day:02d}"

        # Title: nearest preceding non-skip, non-Scripture, non-Date line
        title = ""
        for back in range(date_pos - 1, -1, -1):
            cand = paragraphs[back]
            if cand in skip_titles:
                continue
            if DATE_RE.match(cand):
                continue
            if SCRIPTURE_HEAD_RE.match(cand):
                continue
            # Don't grab a body paragraph from a previous entry
            if back > 0 and back - 1 >= 0 and idx > 0:
                # If this candidate sits BEFORE the previous Date line,
                # it belongs to the previous entry — skip.
                prev_date = date_positions[idx - 1]
                if back <= prev_date:
                    title = ""
                    break
            title = cand
            break

        if not title:
            title = f"Reflection {iso_date}"

        # Body + scripture: from date_pos+1 to (next Date - title line, exclusive)
        body_end = (date_positions[idx + 1] - 1) if idx + 1 < len(date_positions) else len(paragraphs)
        # If next entry has its own title, exclude that title line too
        if idx + 1 < len(date_positions):
            # Find nearest non-skip line before the next date that's a title
            for back in range(date_positions[idx + 1] - 1, date_pos, -1):
                cand = paragraphs[back]
                if cand in skip_titles:
                    continue
                if SCRIPTURE_HEAD_RE.match(cand):
                    continue
                if DATE_RE.match(cand):
                    continue
                body_end = back
                break

        body_parts: list[str] = []
        scripture_ref = ""
        scripture_text_parts: list[str] = []
        in_scripture = False

        for k in range(date_pos + 1, body_end):
            p = paragraphs[k]
            if p in skip_titles:
                continue
            scr = SCRIPTURE_HEAD_RE.match(p)
            if scr:
                in_scripture = True
                scripture_ref = (scr.group(1) or "").strip()
                continue
            if in_scripture:
                scripture_text_parts.append(p)
            else:
                body_parts.append(p)

        body_text = "\n\n".join(b for b in body_parts if b)
        scripture_text = "\n".join(s for s in scripture_text_parts if s)

        entry_id = f"dev_{iso_date}_{_slugify(title, 30)}_{_short_hash(title + iso_date)}"
        entries.append({
            "id": entry_id,
            "kind": "devotional",
            "title": title,
            "date": iso_date,
            "body": body_text,
            "scripture_ref": scripture_ref,
            "scripture_text": scripture_text,
            "source_file": path.name,
            "author": "Matt Harris",
        })

    return entries


def extract_sermon_outline(path: Path) -> dict:
    """Parse the Raven & Dove sermon outline docx into a structured record."""
    d = Document(path)
    paragraphs = [p.text.strip() for p in d.paragraphs]
    paragraphs = [p for p in paragraphs if p]
    title = paragraphs[0] if paragraphs else "Sermon"
    body = "\n\n".join(paragraphs[1:])
    return {
        "id": f"sermon_{_slugify(title, 40)}_{_short_hash(title)}",
        "kind": "sermon",
        "title": title.replace("Sermon Outline:", "").strip(),
        "format": "outline",
        "body": body,
        "source_file": path.name,
        "author": "Matt Harris",
        "primary_scripture": "Genesis 8:6-12",
    }


def extract_sermon_pdf(path: Path, primary_scripture: str) -> dict:
    if PdfReader is None:
        return {}
    r = PdfReader(str(path))
    text = "\n".join(p.extract_text() for p in r.pages).strip()
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    title = lines[0] if lines else path.stem
    return {
        "id": f"sermon_{_slugify(title, 40)}_{_short_hash(title)}",
        "kind": "sermon",
        "title": title,
        "format": "full",
        "body": text,
        "source_file": path.name,
        "author": "Matt Harris",
        "primary_scripture": primary_scripture,
    }


def main():
    DEVOTIONALS_OUT.parent.mkdir(parents=True, exist_ok=True)
    SERMONS_OUT.parent.mkdir(parents=True, exist_ok=True)

    # Devotionals
    all_devotionals: list[dict] = []
    for batch in DEVOTIONAL_BATCHES:
        if not batch.exists():
            print(f"missing: {batch}", file=sys.stderr)
            continue
        entries = extract_devotional_entries(batch)
        all_devotionals.extend(entries)
        print(f"  {batch.name}: {len(entries)} entries")

    all_devotionals.sort(key=lambda e: e["date"])
    with DEVOTIONALS_OUT.open("w", encoding="utf-8") as fh:
        for e in all_devotionals:
            fh.write(json.dumps(e, ensure_ascii=False) + "\n")
    print(f"wrote {len(all_devotionals)} devotional reflections to {DEVOTIONALS_OUT.relative_to(REPO)}")

    # Sermons
    sermons: list[dict] = []
    raven = SOURCE_DIR / "Sermon_Outline_Raven_and_Dove_Full.docx"
    if raven.exists():
        sermons.append(extract_sermon_outline(raven))
    s1 = SOURCE_DIR / "Sermon1_Cornerstone.pdf"
    if s1.exists():
        sermons.append(extract_sermon_pdf(s1, "Psalm 118:21-23"))
    s2 = SOURCE_DIR / "Sermon2_Restoration.pdf"
    if s2.exists():
        sermons.append(extract_sermon_pdf(s2, "1 Peter 5:9-11"))

    with SERMONS_OUT.open("w", encoding="utf-8") as fh:
        for s in sermons:
            fh.write(json.dumps(s, ensure_ascii=False) + "\n")
    print(f"wrote {len(sermons)} sermons to {SERMONS_OUT.relative_to(REPO)}")


if __name__ == "__main__":
    main()
